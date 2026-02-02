import asyncio
import io
import json
import logging
import os
import re
from contextlib import asynccontextmanager
from datetime import datetime, timedelta
from typing import Any, Dict, List, Optional, Tuple

from dotenv import load_dotenv
from fastapi import BackgroundTasks, FastAPI, File, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, HTMLResponse
from PIL import Image, ImageOps

import google.generativeai as genai

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


# ----------------------------
# Logging Setup
# ----------------------------
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S"
)
logger = logging.getLogger(__name__)


# ----------------------------
# Env & Paths
# ----------------------------
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(BASE_DIR)
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
OUTPUT_DIR = os.path.join(BASE_DIR, "output")
FRONT_DIST_DIR = os.path.join(ROOT_DIR, "front", "dist")
FRONT_ASSETS_DIR = os.path.join(FRONT_DIST_DIR, "assets")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

RECEIPT_RETENTION_HOURS = int(os.getenv("RECEIPT_RETENTION_HOURS", "3"))

MODEL_NAME = "gemini-2.5-flash"


# ----------------------------
# Constants
# ----------------------------
MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
IMAGE_MAX_SIDE = 1800
JPEG_QUALITY = 85
FONT_SIZE_NORMAL = 11
FONT_SIZE_TITLE = 18
FONT_SIZE_ITEM_TITLE = 14
CLEANUP_INTERVAL_SECONDS = 600  # 10분


# ----------------------------
# Pre-compiled Regex Patterns
# ----------------------------
JSON_PATTERN = re.compile(r"\{.*\}", re.DOTALL)
DATE_PATTERN = re.compile(r"\d{4}-\d{2}-\d{2}")
SAFE_FILENAME_PATTERN = re.compile(r"[^a-zA-Z0-9가-힣_-]")
CODE_BLOCK_START_PATTERN = re.compile(r"^```(?:json)?\s*", re.IGNORECASE)
CODE_BLOCK_END_PATTERN = re.compile(r"\s*```$")
DIGIT_ONLY_PATTERN = re.compile(r"[^\d]")


# ----------------------------
# Account Info (from env)
# ----------------------------
def load_accounts_from_env() -> Dict[str, Dict[str, str]]:
    """Load account info from environment variable."""
    accounts_json = os.getenv("ACCOUNTS_JSON", "")
    if accounts_json:
        try:
            return json.loads(accounts_json)
        except json.JSONDecodeError:
            logger.warning("Failed to parse ACCOUNTS_JSON, using defaults")
    # Fallback to individual env vars
    return {
        "이진모": {
            "bank": os.getenv("ACCOUNT_JINMO_BANK", ""),
            "account": os.getenv("ACCOUNT_JINMO_NUMBER", "")
        },
        "박지민": {
            "bank": os.getenv("ACCOUNT_JIMIN_BANK", ""),
            "account": os.getenv("ACCOUNT_JIMIN_NUMBER", "")
        },
    }

ACCOUNTS = load_accounts_from_env()


# ----------------------------
# CORS Configuration
# ----------------------------
ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost:8000,http://localhost:5173,http://127.0.0.1:8000,http://127.0.0.1:5173"
).split(",")


# ----------------------------
# Background Cleanup Task
# ----------------------------
async def periodic_cleanup():
    """Background task to periodically clean up expired files."""
    while True:
        try:
            removed = cleanup_expired_files()
            if removed > 0:
                logger.info(f"Cleanup: removed {removed} expired files")
        except Exception as e:
            logger.error(f"Cleanup error: {e}")
        await asyncio.sleep(CLEANUP_INTERVAL_SECONDS)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Lifespan context manager for startup/shutdown events."""
    # Startup: start background cleanup task
    cleanup_task = asyncio.create_task(periodic_cleanup())
    logger.info("Started periodic cleanup task")
    yield
    # Shutdown: cancel cleanup task
    cleanup_task.cancel()
    try:
        await cleanup_task
    except asyncio.CancelledError:
        pass
    logger.info("Stopped periodic cleanup task")


# ----------------------------
# App
# ----------------------------
app = FastAPI(title="CIS Billing Doc Automation", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

if os.path.isdir(FRONT_ASSETS_DIR):
    app.mount("/assets", StaticFiles(directory=FRONT_ASSETS_DIR), name="assets")
if os.path.isdir(UPLOAD_DIR):
    app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")


# ----------------------------
# Utils
# ----------------------------
def now_ts() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def cleanup_expired_files() -> int:
    cutoff = datetime.now() - timedelta(hours=RECEIPT_RETENTION_HOURS)
    removed = 0
    for directory in (UPLOAD_DIR, OUTPUT_DIR):
        if not os.path.isdir(directory):
            continue
        for name in os.listdir(directory):
            path = os.path.join(directory, name)
            if not os.path.isfile(path):
                continue
            try:
                mtime = datetime.fromtimestamp(os.path.getmtime(path))
                if mtime < cutoff:
                    os.remove(path)
                    removed += 1
            except OSError:
                continue
    return removed


def safe_filename(original: str) -> str:
    original = os.path.basename(original)
    original = SAFE_FILENAME_PATTERN.sub("_", original)
    return original[:120]


def image_to_jpeg_bytes(file_bytes: bytes, max_side: int = IMAGE_MAX_SIDE, quality: int = JPEG_QUALITY) -> bytes:
    """Gemini OCR 안정성을 위해 이미지를 RGB JPEG로 변환하고 리사이즈."""
    with Image.open(io.BytesIO(file_bytes)) as im:
        im = ImageOps.exif_transpose(im)
        im = im.convert("RGB")
        w, h = im.size
        scale = min(1.0, max_side / float(max(w, h)))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)))
        out = io.BytesIO()
        im.save(out, format="JPEG", quality=quality, optimize=True)
        return out.getvalue()


def normalize_image_bytes(file_bytes: bytes) -> bytes:
    try:
        with Image.open(io.BytesIO(file_bytes)) as im:
            im = ImageOps.exif_transpose(im)
            fmt = (im.format or "JPEG").upper()
            out = io.BytesIO()
            save_kwargs = {}
            if fmt == "JPEG":
                im = im.convert("RGB")
                save_kwargs = {"quality": 92, "optimize": True}
            im.save(out, format=fmt, **save_kwargs)
            return out.getvalue()
    except Exception:
        return file_bytes


def parse_json_strict(text: str) -> Dict[str, Any]:
    """
    Gemini가 코드블록/주석 등을 섞어 출력하는 경우를 대비해 JSON 객체만 추출.
    """
    if not text:
        raise ValueError("Empty response")

    text = text.strip()
    text = CODE_BLOCK_START_PATTERN.sub("", text)
    text = CODE_BLOCK_END_PATTERN.sub("", text)

    m = JSON_PATTERN.search(text)
    if not m:
        raise ValueError("No JSON object found")
    json_str = m.group(0)

    return json.loads(json_str)


def normalize_date(date_str: Optional[str]) -> Optional[str]:
    if not date_str:
        return None
    date_str = date_str.strip()
    if DATE_PATTERN.fullmatch(date_str):
        return date_str
    return None


def normalize_amount(amount: Any) -> Optional[int]:
    if amount is None:
        return None
    if isinstance(amount, int):
        return amount
    if isinstance(amount, float):
        return int(amount)
    if isinstance(amount, str):
        s = DIGIT_ONLY_PATTERN.sub("", amount)
        if s:
            return int(s)
    return None


def set_korean_font(run):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")


def add_run_k(doc_paragraph, text: str, bold: bool = False, size_pt: Optional[int] = None):
    run = doc_paragraph.add_run(text)
    set_korean_font(run)
    run.bold = bold
    if size_pt:
        run.font.size = Pt(size_pt)
    return run


def won(n: int) -> str:
    return f"{n:,}"


def item_total_amount(item: Dict[str, Any]) -> int:
    receipts = item.get("receipts", []) or []
    receipts_total = sum((r.get("amount") or 0) for r in receipts if isinstance(r.get("amount"), int))
    if receipts_total:
        return receipts_total
    total_amount = item.get("totalAmount")
    return int(total_amount) if isinstance(total_amount, int) else 0


def set_table_borders(table, size_pt: int = 1, color: str = "000000"):
    size = str(int(size_pt * 8))
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.find(qn("w:tblBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = tblBorders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            tblBorders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str = "DDDDDD"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def country_name_ko(code: str) -> str:
    mapping = {
        "RUS": "러시아",
        "CRM": "크림지역",
        "KAZ": "카자흐스탄",
        "UZB": "우즈베키스탄",
        "UKR": "우크라이나",
    }
    return mapping.get(code, "")


def currency_unit_for_item(language: str, item: Dict[str, Any]) -> str:
    if language != "ru":
        return "원"
    code = str(item.get("countryCode") or "").strip().upper()
    mapping = {
        "RUS": "루블",
        "CRM": "루블",
        "KAZ": "텡게",
        "UZB": "숨",
        "UKR": "흐리우냐",
    }
    return mapping.get(code, "루블")


def add_paragraph_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = pBdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        pBdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "000000")


def receipt_image_width(receipt_count: int) -> float:
    rows = (receipt_count + 1) // 2
    if rows <= 1:
        return 3.2
    if rows == 2:
        return 2.6
    if rows == 3:
        return 2.1
    if rows == 4:
        return 1.7
    return 1.4


def receipt_image_max_dimensions(receipt_count: int) -> Tuple[float, float]:
    cols = 2
    rows = (receipt_count + cols - 1) // cols
    max_width = receipt_image_width(receipt_count)
    max_height_total = 6.3
    max_height = max_height_total / max(1, rows)
    return max_width, max_height


def fit_image_dimensions(jpeg_bytes: bytes, max_width: float, max_height: float) -> Tuple[float, float]:
    with Image.open(io.BytesIO(jpeg_bytes)) as im:
        w, h = im.size
    if w <= 0 or h <= 0:
        return max_width, max_height
    if max_width <= 0 or max_height <= 0:
        return max_width, max_height
    aspect = w / h
    max_aspect = max_width / max_height
    if aspect >= max_aspect:
        width = max_width
        height = max_width / aspect
    else:
        height = max_height
        width = max_height * aspect
    return width, height


# ----------------------------
# Gemini OCR
# ----------------------------
OCR_PROMPT = """영수증 이미지를 분석해서 아래 정보를 추출해줘.
1) 날짜 (YYYY-MM-DD)
2) 금액 (KRW만 추출, 다른 통화는 무시)
3) 다중 통화 포함 여부
4) 영수증 원문 텍스트 (rawText)
5) 상호명/가맹점명 (merchant)

응답은 반드시 JSON 형식으로만 반환:
{
  "date": "2025-01-20",
  "amount": 150000,
  "hasMultipleCurrency": false,
  "rawText": "영수증 원문 텍스트",
  "merchant": "상호명"
}

중요:
- 금액은 숫자만 (기호/쉼표 제거)
- 날짜를 찾지 못하면 null
- 다른 통화가 함께 있으면 hasMultipleCurrency를 true
"""

TRANSLATE_PROMPT = """Translate the following JSON to Korean.
Return ONLY a JSON object with the same structure and keys.
Rules:
- Preserve numbers, dates, ids, file names, and account numbers as-is.
- Do NOT translate recipientType values.
- Do NOT translate the names "이진모" and "박지민".
- Translate only human-language fields (description, recipient, bank, rawText, managerName).
- For managerName, use a natural Korean phonetic transcription of the original name.
"""

TRANSLATE_PROMPT_NATURAL = """Translate the following JSON to Korean.
Return ONLY a JSON object with the same structure and keys.
Rules:
- Preserve numbers, dates, ids, file names, and account numbers as-is.
- Do NOT translate recipientType values.
- Do NOT translate the names "이진모" and "박지민".
- Translate only human-language fields (description, recipient, bank, rawText, managerName).
- For managerName, use a natural Korean phonetic transcription of the original name.
- Make the Korean natural and business-appropriate; avoid literal translation.
"""


def ensure_gemini_ready() -> bool:
    return bool(GEMINI_API_KEY)


def doc_country_code(items: List[Dict[str, Any]]) -> str:
    if not items:
        return ""
    first = items[0] if isinstance(items[0], dict) else {}
    return str(first.get("countryCode") or "").strip().upper()


def record_output_manifest(output_name: str, items: List[Dict[str, Any]]):
    receipt_files: List[str] = []
    for it in items:
        receipts = it.get("receipts") or []
        for r in receipts:
            name = str(r.get("fileName") or "").strip()
            if name:
                receipt_files.append(os.path.basename(name))
    data = {
        "receipts": sorted(set(receipt_files)),
        "createdAt": datetime.now().isoformat(),
    }
    manifest_path = os.path.join(OUTPUT_DIR, f"{output_name}.json")
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)


def delete_files(paths: List[str]):
    for path in paths:
        try:
            if os.path.isfile(path):
                os.remove(path)
        except OSError:
            continue


def translate_items_to_korean(items: List[Dict[str, Any]], natural: bool = False) -> List[Dict[str, Any]]:
    if not items:
        return items
    if not ensure_gemini_ready():
        raise RuntimeError("GEMINI_API_KEY not set")

    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(MODEL_NAME)
    payload = {"items": items}
    prompt_base = TRANSLATE_PROMPT_NATURAL if natural else TRANSLATE_PROMPT
    prompt = f"{prompt_base}\nINPUT:\n{json.dumps(payload, ensure_ascii=False)}"
    resp = model.generate_content(prompt)
    data = parse_json_strict(resp.text or "")
    translated = data.get("items")
    if not isinstance(translated, list):
        raise ValueError("Translation response missing items")
    return translated


def gemini_ocr(file_bytes: bytes) -> Dict[str, Any]:
    if not ensure_gemini_ready():
        raise RuntimeError("GEMINI_API_KEY not set")

    genai.configure(api_key=GEMINI_API_KEY)
    model = genai.GenerativeModel(MODEL_NAME)

    jpeg_bytes = image_to_jpeg_bytes(file_bytes)
    img_part = {"mime_type": "image/jpeg", "data": jpeg_bytes}

    resp = model.generate_content([OCR_PROMPT, img_part])
    text = (resp.text or "").strip()

    data = parse_json_strict(text)
    date_norm = normalize_date(data.get("date"))
    amount_norm = normalize_amount(data.get("amount"))

    merchant = str(data.get("merchant") or "").strip()
    return {
        "date": date_norm,
        "amount": amount_norm,
        "hasMultipleCurrency": bool(data.get("hasMultipleCurrency", False)),
        "rawText": data.get("rawText", "") or "",
        "merchant": merchant,
    }


# ----------------------------
# DOCX Generation
# ----------------------------
def docx_set_default_style(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(FONT_SIZE_NORMAL)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    para = style.paragraph_format
    para.space_before = Pt(0)
    para.space_after = Pt(0)
    para.line_spacing = 1.0


def set_row_height(row, cm_value: float = 0.65):
    row.height = Cm(cm_value)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def add_summary_table(
    doc: Document,
    items: List[Dict[str, Any]],
    language: str,
    report_year: Optional[int] = None,
    report_month: Optional[int] = None,
):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_suffix = ""
    if report_year and report_month:
        report_suffix = f" - {report_year % 100}년 {int(report_month)}월"
    if language == "ru":
        country = country_name_ko(doc_country_code(items))
        title = f"CIS 지역 - 지출내역서 - {country}{report_suffix}".strip()
    else:
        title = f"CIS 청구 문서{report_suffix}"
    add_run_k(title_p, title, bold=True, size_pt=FONT_SIZE_TITLE)
    title_p.paragraph_format.space_after = Pt(12)

    table = doc.add_table(rows=1, cols=5)
    set_table_borders(table, size_pt=1, color="000000")
    hdr = table.rows[0].cells
    if language == "ru":
        headers = ["번호", "내용", "금액", "담당자", "소계"]
    else:
        headers = ["번호", "내용", "금액", "수령인", "소계"]
    for i, h in enumerate(headers):
        para = hdr[i].paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_k(para, h, bold=True)
    set_row_height(table.rows[0])
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "EEEEEE")

    recipients_order: List[str] = []
    grouped: Dict[str, List[Dict[str, Any]]] = {}
    for it in items:
        if language == "ru":
            recipient = str(it.get("managerName") or "").strip()
        else:
            recipient = str(it.get("recipient", it.get("recipientType", "")))
        recipient = recipient or "-"
        if recipient not in grouped:
            grouped[recipient] = []
            recipients_order.append(recipient)
        grouped[recipient].append(
            {
                "description": str(it.get("description", "")),
                "amount": item_total_amount(it),
            }
        )

    total = 0
    row_index = 1
    for recipient in recipients_order:
        rows = grouped[recipient]
        group_total = sum((r.get("amount") or 0) for r in rows if isinstance(r.get("amount"), int))
        total += group_total
        group_start = row_index
        for r in rows:
            row = table.add_row().cells
            set_row_height(table.rows[row_index])
            amount = r.get("amount")
            if isinstance(amount, int):
                unit = currency_unit_for_item(language, items[0]) if items else ""
                amount_text = f"{won(amount)} {unit}".strip() if language == "ru" else won(amount)
            else:
                amount_text = ""
            show_group = row_index == group_start
            row_vals = [
                str(row_index),
                str(r.get("description", "")),
                amount_text,
                recipient if show_group else "",
                (
                    f"{won(group_total)} {currency_unit_for_item(language, items[0])}".strip()
                    if (show_group and language == "ru")
                    else (won(group_total) if show_group else "")
                ),
            ]
            for i, v in enumerate(row_vals):
                para = row[i].paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i in (0, 2, 3, 4) else WD_ALIGN_PARAGRAPH.LEFT
                add_run_k(para, v)
            row_index += 1
        if len(rows) > 1:
            table.cell(group_start, 3).merge(table.cell(row_index - 1, 3))
            table.cell(group_start, 4).merge(table.cell(row_index - 1, 4))

    sum_row = table.add_row().cells
    set_row_height(table.rows[row_index])
    sum_row[0].merge(sum_row[1]).merge(sum_row[2]).merge(sum_row[3])
    para = sum_row[0].paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_k(para, "합계", bold=True)
    para2 = sum_row[4].paragraphs[0]
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_text = (
        f"{won(total)} {currency_unit_for_item(language, items[0])}".strip()
        if language == "ru"
        else won(total)
    )
    add_run_k(para2, total_text, bold=True)
    for cell in table.rows[row_index].cells:
        set_cell_shading(cell, "EEEEEE")

    col_widths = [Inches(0.5), Inches(3.4), Inches(1.0), Inches(1.0), Inches(1.1)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]


def add_item_detail(doc: Document, item: Dict[str, Any], item_index: int, language: str):
    doc.add_page_break()
    title_p = doc.add_paragraph()
    title_text = str(item.get("description") or "").strip() or "상세내용"
    add_run_k(title_p, f"{item_index}. {title_text}", bold=True, size_pt=FONT_SIZE_ITEM_TITLE)

    info = doc.add_paragraph()
    add_run_k(info, "영수증 일자: ", bold=True)
    add_run_k(info, str(item.get("date", "")))

    receipts: List[Dict[str, Any]] = item.get("receipts", []) or []
    receipts_total = item_total_amount(item)
    unit = currency_unit_for_item(language, item)

    amt_p = doc.add_paragraph()
    add_run_k(amt_p, "금액: ", bold=True)
    amount_text = f"{won(receipts_total)}{unit}" if language != "ru" else f"{won(receipts_total)} {unit}"
    add_run_k(amt_p, amount_text, bold=True)

    if language == "ru":
        manager = str(item.get("managerName") or "").strip()
        telegram_id = str(item.get("telegramId") or "").strip()
        acc_p = doc.add_paragraph()
        add_run_k(acc_p, "담당자: ", bold=True)
        add_run_k(acc_p, manager)
        acc_p2 = doc.add_paragraph()
        add_run_k(acc_p2, "텔레그램 ID: ", bold=True)
        add_run_k(acc_p2, telegram_id)
    else:
        account_line = f"{item.get('bank', '')} / {item.get('account', '')} / {item.get('recipient', item.get('recipientType', ''))}"
        acc_p = doc.add_paragraph()
        add_run_k(acc_p, "계좌정보: ", bold=True)
        add_run_k(acc_p, account_line.strip())

    sep_p = doc.add_paragraph()
    add_paragraph_bottom_border(sep_p)

    if not receipts:
        add_run_k(doc.add_paragraph(), "영수증 없음")
        return

    cols = 2
    rows = (len(receipts) + cols - 1) // cols
    grid = doc.add_table(rows=rows, cols=cols)

    img_max_width, img_max_height = receipt_image_max_dimensions(len(receipts))
    for i, r in enumerate(receipts):
        rr = i // cols
        cc = i % cols
        cell = grid.cell(rr, cc)

        amt = r.get("amount")
        if isinstance(amt, int) and amt != 0:
            label = f"영수증{i+1} - {won(amt)}원"
        else:
            label = f"영수증{i+1}"
        p1 = cell.paragraphs[0]
        add_run_k(p1, label, bold=True)

        file_name = r.get("fileName")
        if file_name:
            img_path = os.path.join(UPLOAD_DIR, os.path.basename(file_name))
            if os.path.exists(img_path):
                cell.add_paragraph("")
                try:
                    run = cell.add_paragraph().add_run()
                    set_korean_font(run)
                    with open(img_path, "rb") as f:
                        jpeg_bytes = image_to_jpeg_bytes(f.read())
                    width_in, height_in = fit_image_dimensions(
                        jpeg_bytes, img_max_width, img_max_height
                    )
                    run.add_picture(
                        io.BytesIO(jpeg_bytes),
                        width=Inches(width_in),
                        height=Inches(height_in),
                    )
                except Exception:
                    errp = cell.add_paragraph()
                    add_run_k(errp, "(이미지 삽입 실패)")

    if len(receipts) % cols == 1:
        last_cell = grid.cell(rows - 1, 1)
        last_cell.text = ""


def validate_items(items: List[Dict[str, Any]], language: str) -> List[str]:
    errors: List[str] = []
    for idx, it in enumerate(items, start=1):
        prefix = f"[항목 {idx}] "
        if not (it.get("description") or "").strip():
            errors.append(prefix + "상세 내용이 비어 있습니다.")
        if not (it.get("date") or "").strip():
            errors.append(prefix + "결제일자가 필요합니다.")
        receipts = it.get("receipts") or []
        if len(receipts) < 1:
            errors.append(prefix + "영수증이 최소 1개 필요합니다.")
        else:
            first = receipts[0]
            if not (first.get("date") or "").strip():
                errors.append(prefix + "첫 번째 영수증 날짜가 필요합니다.")
            if first.get("amount") in (None, "", 0):
                errors.append(prefix + "첫 번째 영수증 금액이 필요합니다.")

        if language != "ru":
            rtype = it.get("recipientType")
            if rtype not in ("이진모", "박지민", "기타"):
                errors.append(prefix + "수령인 구분이 필요합니다.")
            if rtype == "기타":
                if not (it.get("recipient") or "").strip():
                    errors.append(prefix + "기타 수령인명이 필요합니다.")
                if not (it.get("bank") or "").strip():
                    errors.append(prefix + "기타 은행명이 필요합니다.")
                if not (it.get("account") or "").strip():
                    errors.append(prefix + "기타 계좌번호가 필요합니다.")
    return errors


# ----------------------------
# Routes
# ----------------------------
@app.get("/api/accounts")
def get_accounts():
    return {
        "accounts": ACCOUNTS,
        "server_info": {
            "model": MODEL_NAME,
            "status": "연결됨" if ensure_gemini_ready() else "API 키 필요",
            "api_ready": ensure_gemini_ready(),
        },
    }


@app.get("/downloads/{filename}")
def download_file(filename: str, background_tasks: BackgroundTasks):
    safe_name = os.path.basename(filename)
    file_path = os.path.join(OUTPUT_DIR, safe_name)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="file not found")

    manifest_path = os.path.join(OUTPUT_DIR, f"{safe_name}.json")
    delete_paths = [file_path, manifest_path]
    if os.path.exists(manifest_path):
        try:
            with open(manifest_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            for name in data.get("receipts", []):
                if name:
                    delete_paths.append(os.path.join(UPLOAD_DIR, os.path.basename(name)))
        except (OSError, json.JSONDecodeError):
            pass

    background_tasks.add_task(delete_files, delete_paths)
    return FileResponse(file_path, filename=safe_name, background=background_tasks)


@app.get("/", include_in_schema=False)
def serve_root():
    index_path = os.path.join(FRONT_DIST_DIR, "index.html")
    if not os.path.exists(index_path):
        return HTMLResponse("frontend build not found", status_code=404)
    return FileResponse(index_path)


@app.get("/{full_path:path}", include_in_schema=False)
def serve_spa(full_path: str):
    if full_path.startswith(("api/", "downloads/", "assets/")):
        raise HTTPException(status_code=404, detail="Not Found")
    index_path = os.path.join(FRONT_DIST_DIR, "index.html")
    if not os.path.exists(index_path):
        return HTMLResponse("frontend build not found", status_code=404)
    return FileResponse(index_path)


@app.post("/api/upload-receipt")
async def upload_receipt(file: UploadFile = File(...)):
    if not file:
        raise HTTPException(status_code=400, detail="No file")

    cleanup_expired_files()

    content = await file.read()
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=400, detail=f"File too large (max {MAX_FILE_SIZE_MB}MB)")

    normalized_content = normalize_image_bytes(content)
    fname = f"{now_ts()}_{safe_filename(file.filename or 'receipt')}"
    save_path = os.path.join(UPLOAD_DIR, fname)

    with open(save_path, "wb") as f:
        f.write(normalized_content)

    ocr_payload: Dict[str, Any] = {
        "date": None,
        "amount": None,
        "rawText": "",
        "merchant": "",
        "hasMultipleCurrency": False,
    }
    ocr_success = False
    ocr_error = ""
    if ensure_gemini_ready():
        try:
            ocr_payload = gemini_ocr(normalized_content)
            ocr_success = True
            logger.info(
                "[OCR][upload] date=%s amount=%s merchant=%s",
                ocr_payload.get("date"),
                ocr_payload.get("amount"),
                ocr_payload.get("merchant"),
            )
        except Exception as e:
            ocr_error = str(e)
            logger.error("[OCR][upload] failed: %s", ocr_error)

    return {
        "success": True,
        "fileName": fname,
        "date": ocr_payload.get("date"),
        "amount": ocr_payload.get("amount"),
        "rawText": ocr_payload.get("rawText"),
        "merchant": ocr_payload.get("merchant"),
        "hasMultipleCurrency": ocr_payload.get("hasMultipleCurrency"),
        "ocrSuccess": ocr_success,
        "ocrError": ocr_error,
    }


@app.post("/api/ocr")
async def ocr(file: UploadFile = File(...)):
    if not file:
        raise HTTPException(status_code=400, detail="No file")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE_BYTES:
        raise HTTPException(status_code=400, detail=f"File too large (max {MAX_FILE_SIZE_MB}MB)")

    logger.info("[OCR] request started: %s", file.filename)

    try:
        normalized_content = normalize_image_bytes(content)
        result = gemini_ocr(normalized_content)
        logger.info("[OCR] success: date=%s amount=%s", result.get("date"), result.get("amount"))
        return {"success": True, **result}
    except Exception as e:
        logger.error("[OCR] failed: %s", repr(e))
        return {
            "success": False,
            "error": str(e),
            "date": None,
            "amount": None,
            "hasMultipleCurrency": False,
            "rawText": "",
            "merchant": "",
        }


@app.post("/api/ocr-from-upload")
async def ocr_from_upload(payload: Dict[str, Any]):
    file_name = payload.get("fileName")
    if not file_name:
        raise HTTPException(status_code=400, detail="fileName is required")

    img_path = os.path.join(UPLOAD_DIR, os.path.basename(str(file_name)))
    if not os.path.exists(img_path):
        raise HTTPException(status_code=404, detail="file not found")

    try:
        with open(img_path, "rb") as f:
            content = f.read()
        result = gemini_ocr(content)
        logger.info("[OCR][from-upload] date=%s amount=%s", result.get("date"), result.get("amount"))
        return {"success": True, **result}
    except Exception as e:
        logger.error("[OCR][from-upload] failed: %s", str(e))
        return {
            "success": False,
            "error": str(e),
            "date": None,
            "amount": None,
            "hasMultipleCurrency": False,
            "rawText": "",
            "merchant": "",
        }


@app.post("/api/generate-document")
async def generate_document(payload: Dict[str, Any]):
    items = payload.get("items", [])
    if not isinstance(items, list) or not items:
        raise HTTPException(status_code=400, detail="items is required")

    language = str(payload.get("language") or "ko").lower()
    natural_translation = bool(payload.get("naturalTranslation"))
    report_year = payload.get("reportYear")
    report_month = payload.get("reportMonth")
    try:
        report_year = int(report_year) if report_year is not None else None
    except (TypeError, ValueError):
        report_year = None
    try:
        report_month = int(report_month) if report_month is not None else None
    except (TypeError, ValueError):
        report_month = None

    logger.info("[Document] generating: language=%s, items=%d", language, len(items))

    if language == "ru":
        try:
            items = translate_items_to_korean(items, natural=natural_translation)
        except Exception as e:
            logger.error("[Document] translation failed: %s", e)
            raise HTTPException(status_code=400, detail=f"Translation failed: {e}")

    errors = validate_items(items, language)
    if errors:
        logger.warning("[Document] validation failed: %s", errors)
        raise HTTPException(status_code=400, detail={"message": "Validation failed", "errors": errors})

    try:
        doc = Document()
        docx_set_default_style(doc)

        section = doc.sections[0]
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

        add_summary_table(doc, items, language, report_year, report_month)

        for i, it in enumerate(items, start=1):
            add_item_detail(doc, it, i, language)

        if language == "ru":
            country_code = doc_country_code(items)
            code_map = {
                "RUS": "RUS",
                "CRM": "CRM",
                "KAZ": "KAZ",
                "UZB": "UZB",
                "UKR": "UKR",
            }
            code = code_map.get(country_code, "RUS")
            base_name = f"CIS-Recipe-{code}-{datetime.now().strftime('%Y%m%d')}"
            out_name = f"{base_name}.docx"
            if os.path.exists(os.path.join(OUTPUT_DIR, out_name)):
                idx = 1
                while True:
                    candidate = f"{base_name}_{idx}.docx"
                    if not os.path.exists(os.path.join(OUTPUT_DIR, candidate)):
                        out_name = candidate
                        break
                    idx += 1
        else:
            out_name = f"CIS-recipe_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        out_path = os.path.join(OUTPUT_DIR, out_name)
        doc.save(out_path)
        record_output_manifest(out_name, items)

        logger.info("[Document] generated successfully: %s", out_name)
        return {"success": True, "downloadUrl": f"/downloads/{out_name}"}

    except Exception as e:
        logger.error("[Document] generation failed: %s", repr(e))
        raise HTTPException(status_code=500, detail=f"Document generation failed: {str(e)}")


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
